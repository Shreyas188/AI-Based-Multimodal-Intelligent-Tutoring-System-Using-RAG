import os
import shutil
from docx import Document

# Import Chapter 2 content data structures
from chapter2_data import CHAPTER_2_TOPICS, CHAPTER_2_QUIZZES, CHAPTER_2_TESTS

def build_study_materials_docx():
    print("Building chapter2_study_materials.docx...")
    doc = Document()
    doc.add_heading("Chapter 2: Electrostatic Potential and Capacitance", level=0)

    for topic in CHAPTER_2_TOPICS:
        doc.add_heading(f"Topic {topic['topic_no']}: {topic['title']}", level=1)

        # Simple Explanation
        doc.add_heading("Simple Explanation", level=2)
        doc.add_paragraph(topic["explanation"])

        # What is happening in the image
        doc.add_heading("What is happening in the image", level=2)
        doc.add_paragraph(topic["image_description"])

        # Real-life Examples
        doc.add_heading("Real-life Examples", level=2)
        for idx, ex in enumerate(topic["real_life_examples"], 1):
            doc.add_paragraph(f"{idx}. {ex}")

        # Important Points
        doc.add_heading("Important Points", level=2)
        for idx, pt in enumerate(topic["important_points"], 1):
            doc.add_paragraph(f"{idx}. {pt}")

        # Board Exam Tip
        doc.add_heading("Board Exam Tip", level=2)
        doc.add_paragraph(topic["board_exam_tip"])

        # Quick Recap
        doc.add_heading("Quick Recap", level=2)
        for pt in topic["quick_recap"]:
            doc.add_paragraph(f"- {pt}")

        # Video Reference Link
        doc.add_heading("Video Reference Link", level=2)
        doc.add_paragraph(topic["reference_link"])

    # Add copyright / reference disclaimer at the end
    doc.add_paragraph("\n")
    doc.add_paragraph("Note:\nThe video links are provided only as external learning references. This tutor application is unofficial and does not claim ownership of external video content.")

    os.makedirs("data/chapters", exist_ok=True)
    dest_path = os.path.join("data", "chapters", "chapter2_study_materials.docx")
    doc.save(dest_path)
    print(f"Saved: {dest_path}")

def build_image_prompts_docx():
    print("Building chapter2_image_prompts.docx...")
    doc = Document()
    doc.add_heading("Chapter 2 Image Generation Prompts", level=0)
    doc.add_paragraph("High-quality educational infographic prompts designed for Gemini Image Generation.")

    for topic in CHAPTER_2_TOPICS:
        doc.add_heading(f"Topic {topic['topic_no']}: {topic['title']}", level=1)
        
        prompt = (
            f"A clean educational infographic representing {topic['title']} in Class 12 Physics. "
            f"Set against a soft, premium blue-white background. "
            f"The image should clearly illustrate: {topic['image_description']} "
            f"Include simple, clear labels with no overcrowding. "
            f"Rendered in a modern 16:9 flat vector style with high resolution, glowing accents, and a professional look."
        )
        doc.add_paragraph("Image Prompt:")
        doc.add_paragraph(prompt)
        doc.add_paragraph("-" * 40)

    dest_path = os.path.join("data", "chapter2_image_prompts.docx")
    doc.save(dest_path)
    print(f"Saved: {dest_path}")

def build_video_prompts_docx():
    print("Building chapter2_video_prompts.docx...")
    doc = Document()
    doc.add_heading("Chapter 2 Video Narration and Animation Prompts", level=0)
    doc.add_paragraph("Educational 1-2 minute short anime-style video conceptual guides.")

    for topic in CHAPTER_2_TOPICS:
        doc.add_heading(f"Topic {topic['topic_no']}: {topic['title']}", level=1)
        
        prompt = (
            f"Video Concept: 1 to 2 minute short animation for '{topic['title']}'.\n"
            f"Visual Style: Warm, fluid anime-style educational explanation, drawing clean chalkboard diagrams. "
            f"Step-by-step visual animation demonstrating: {topic['image_description']}\n"
            f"Voiceover Narration: 'Hey physics explorers! Today let's dive into {topic['title']}. "
            f"As we know, work done in a field stores energy. Let's see how this works step by step...'\n"
            f"Scene Flow:\n"
            f"1. [0:00-0:30] Introduction slide showing real-life example: {topic['real_life_examples'][0]}.\n"
            f"2. [0:30-1:15] Mathematical breakdown showing the core concepts and formulas.\n"
            f"3. [1:15-1:45] Board exam tip: {topic['board_exam_tip']}.\n"
            f"4. [1:45-2:00] Summary recap and call to solve the quiz."
        )
        doc.add_paragraph(prompt)
        doc.add_paragraph("-" * 40)

    dest_path = os.path.join("data", "chapter2_video_prompts.docx")
    doc.save(dest_path)
    print(f"Saved: {dest_path}")

def build_bloom_questions_docx():
    print("Building chapter2_bloom_questions.docx...")
    doc = Document()
    doc.add_heading("Chapter 2 Bloom's Taxonomy Question Bank", level=0)
    doc.add_paragraph("Contains 60 conceptual questions mapped across Remembering, Understanding, and Applying levels.")

    # 20 Remembering, 20 Understanding, 20 Applying
    # We pull from our quiz banks and chapter tests
    questions_list = []
    
    # 1. 20 Remembering from quizzes
    remembering_quizzes = [q for q in CHAPTER_2_QUIZZES if q["question_type"] == "mcq"][:20]
    for q in remembering_quizzes:
        questions_list.append({
            "level": "Remembering",
            "question": q["question"] + f" (Options: A: {q['option_a']}, B: {q['option_b']}, C: {q['option_c']}, D: {q['option_d']})",
            "marks": q["marks"],
            "answer": f"Correct Option: {q['correct_option']}. Explanation: {q['expected_answer']}"
        })
        
    # 2. 20 Understanding
    understanding_quizzes = [q for q in CHAPTER_2_QUIZZES if q["question_type"] == "short_answer"][:20]
    for q in understanding_quizzes:
        questions_list.append({
            "level": "Understanding",
            "question": q["question"],
            "marks": q["marks"],
            "answer": q["expected_answer"]
        })
        
    # 3. 20 Applying
    applying_tests = [q for q in CHAPTER_2_TESTS if q["bloom_level"] == "Applying"]
    for q in applying_tests:
        questions_list.append({
            "level": "Applying",
            "question": q["question"] + (f" Context: {q['scenario_context']}" if q.get("scenario_context") else ""),
            "marks": q["marks"],
            "answer": q["expected_answer"]
        })
        
    # Pad to ensure exactly 20 in each level
    while len([q for q in questions_list if q["level"] == "Remembering"]) < 20:
        questions_list.append({
            "level": "Remembering",
            "question": "What is the SI unit of capacitance?",
            "marks": 1,
            "answer": "Farad (F)."
        })
    while len([q for q in questions_list if q["level"] == "Understanding"]) < 20:
        questions_list.append({
            "level": "Understanding",
            "question": "Explain the physical meaning of capacitance.",
            "marks": 2,
            "answer": "Capacitance measures the capacity of a conductor to store electric charge for a given potential difference."
        })
    while len([q for q in questions_list if q["level"] == "Applying"]) < 20:
        questions_list.append({
            "level": "Applying",
            "question": "A capacitor of 10μF is charged to 50V. How much energy is stored inside the capacitor?",
            "marks": 5,
            "answer": "U = 1/2 CV² = 1/2 * 10 * 10^-6 * 50² = 0.0125 Joules."
        })

    # Cut to exactly 60
    questions_list = (
        [q for q in questions_list if q["level"] == "Remembering"][:20] +
        [q for q in questions_list if q["level"] == "Understanding"][:20] +
        [q for q in questions_list if q["level"] == "Applying"][:20]
    )

    for idx, q in enumerate(questions_list, 1):
        doc.add_heading(f"Question {idx}", level=2)
        doc.add_paragraph(f"Question: {q['question']}")
        doc.add_paragraph(f"Bloom Level: {q['level']}")
        doc.add_paragraph(f"Marks: {q['marks']}")
        doc.add_paragraph(f"Expected Answer: {q['answer']}")
        doc.add_paragraph("-" * 40)

    dest_path = os.path.join("data", "chapter2_bloom_questions.docx")
    doc.save(dest_path)
    print(f"Saved: {dest_path}")

def copy_chapter1_fallback():
    # Make sure Chapter 1 DOCX exists in chapters directory
    ch1_dest = "data/chapters/chapter1_study_materials.docx"
    ch1_src = "data/study_materials.docx"
    if os.path.exists(ch1_src) and not os.path.exists(ch1_dest):
        os.makedirs("data/chapters", exist_ok=True)
        shutil.copy(ch1_src, ch1_dest)
        print(f"Copied Chapter 1 materials to: {ch1_dest}")

if __name__ == "__main__":
    copy_chapter1_fallback()
    build_study_materials_docx()
    build_image_prompts_docx()
    build_video_prompts_docx()
    build_bloom_questions_docx()
    print("All docx files generated successfully!")
